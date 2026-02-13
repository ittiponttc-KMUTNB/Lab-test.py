import streamlit as st
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import matplotlib
import io
import json
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

matplotlib.rcParams['font.family'] = 'DejaVu Sans'

# =============================================================
# Page Config
# =============================================================
st.set_page_config(
    page_title="Particleboard Bending Test",
    page_icon="🧪",
    layout="wide",
    initial_sidebar_state="expanded"
)

# =============================================================
# Custom CSS — Clean Engineering Theme
# =============================================================
st.markdown("""
<style>
    /* Main background */
    .stApp {
        background: linear-gradient(135deg, #f8f9fc 0%, #eef1f8 100%);
    }

    /* Sidebar */
    section[data-testid="stSidebar"] {
        background: linear-gradient(180deg, #1e293b 0%, #0f172a 100%);
    }
    section[data-testid="stSidebar"] * {
        color: #e2e8f0 !important;
    }
    section[data-testid="stSidebar"] .stMarkdown h1,
    section[data-testid="stSidebar"] .stMarkdown h2,
    section[data-testid="stSidebar"] .stMarkdown h3 {
        color: #38bdf8 !important;
    }

    /* Tab styling */
    .stTabs [data-baseweb="tab-list"] {
        gap: 8px;
        background: white;
        border-radius: 12px;
        padding: 6px;
        box-shadow: 0 1px 3px rgba(0,0,0,0.06);
    }
    .stTabs [data-baseweb="tab"] {
        border-radius: 8px;
        padding: 10px 24px;
        font-weight: 600;
        font-size: 15px;
    }
    .stTabs [aria-selected="true"] {
        background: linear-gradient(135deg, #2563eb, #3b82f6) !important;
        color: white !important;
        border-radius: 8px;
    }

    /* Cards */
    .info-card {
        background: white;
        border-radius: 14px;
        padding: 20px 24px;
        box-shadow: 0 2px 8px rgba(0,0,0,0.05);
        border-left: 4px solid #3b82f6;
        margin-bottom: 16px;
    }
    .result-card {
        background: linear-gradient(135deg, #ecfdf5 0%, #d1fae5 100%);
        border-radius: 14px;
        padding: 20px 24px;
        border-left: 4px solid #10b981;
        margin-bottom: 12px;
    }
    .warn-card {
        background: linear-gradient(135deg, #fefce8 0%, #fef9c3 100%);
        border-radius: 14px;
        padding: 16px 20px;
        border-left: 4px solid #f59e0b;
        margin-bottom: 12px;
        font-size: 14px;
    }
    .summary-card {
        background: linear-gradient(135deg, #eff6ff 0%, #dbeafe 100%);
        border-radius: 14px;
        padding: 20px 24px;
        border-left: 4px solid #2563eb;
        margin-bottom: 12px;
    }

    /* Metric boxes */
    .metric-box {
        background: white;
        border-radius: 12px;
        padding: 18px;
        text-align: center;
        box-shadow: 0 2px 6px rgba(0,0,0,0.04);
        border: 1px solid #e2e8f0;
    }
    .metric-box .label {
        font-size: 13px;
        color: #64748b;
        font-weight: 500;
        margin-bottom: 4px;
    }
    .metric-box .value {
        font-size: 26px;
        font-weight: 700;
        color: #1e293b;
    }
    .metric-box .unit {
        font-size: 13px;
        color: #94a3b8;
    }

    /* Header banner */
    .header-banner {
        background: linear-gradient(135deg, #1e3a5f 0%, #2563eb 50%, #3b82f6 100%);
        color: white;
        padding: 28px 32px;
        border-radius: 16px;
        margin-bottom: 24px;
        box-shadow: 0 4px 16px rgba(37,99,235,0.2);
    }
    .header-banner h1 {
        margin: 0; font-size: 28px; font-weight: 700;
    }
    .header-banner p {
        margin: 6px 0 0 0; font-size: 15px; opacity: 0.88;
    }

    /* Footer */
    .footer {
        text-align: center;
        padding: 20px;
        margin-top: 32px;
        color: #64748b;
        font-size: 14px;
        border-top: 1px solid #e2e8f0;
    }

    /* Number input compact */
    .stNumberInput > div { max-width: 100%; }

    /* Buttons */
    .stButton > button {
        border-radius: 10px;
        font-weight: 600;
        padding: 8px 20px;
        transition: all 0.2s;
    }
</style>
""", unsafe_allow_html=True)

# =============================================================
# Calculation Functions
# =============================================================

def calc_MOR(Fmax_N: float, L: float, b: float, t: float) -> float:
    """MOR = 3·F·L / (2·b·t²)  [Three-point bending]"""
    if b <= 0 or t <= 0 or L <= 0:
        raise ValueError("L, b, t ต้องมากกว่าศูนย์")
    return (3.0 * Fmax_N * L) / (2.0 * b * t**2)


def calc_slope_elastic(load_N: np.ndarray, defl_mm: np.ndarray,
                       idx_lo: int, idx_hi: int):
    """
    Slope (N/mm) จาก linear fit ของช่วงที่เลือก (idx_lo ถึง idx_hi)
    คืนค่า: (slope, intercept, r_squared)
    """
    x = defl_mm[idx_lo:idx_hi+1]
    y = load_N[idx_lo:idx_hi+1]

    if hasattr(x, 'values'):
        x = x.values
    if hasattr(y, 'values'):
        y = y.values

    if len(x) < 2:
        raise ValueError("ต้องเลือกอย่างน้อย 2 จุดสำหรับ linear fit")

    coeffs = np.polyfit(x, y, 1)
    slope = coeffs[0]      # N/mm
    intercept = coeffs[1]  # N

    y_pred = np.polyval(coeffs, x)
    ss_res = np.sum((y - y_pred)**2)
    ss_tot = np.sum((y - np.mean(y))**2)
    r_sq = 1.0 - ss_res / ss_tot if ss_tot > 0 else 0.0

    return slope, intercept, r_sq


def slope_to_MOE(slope: float, L: float, b: float, t: float) -> float:
    """MOE = slope · L³ / (4·b·t³)  [MPa]"""
    if b <= 0 or t <= 0 or L <= 0:
        raise ValueError("L, b, t ต้องมากกว่าศูนย์")
    return slope * L**3 / (4.0 * b * t**3)


def calc_TS(before: list, after: list) -> tuple:
    """คืนค่า (avg_before, avg_after, ts_percent)"""
    avg_b = np.mean(before)
    avg_a = np.mean(after)
    if avg_b <= 0:
        raise ValueError("ค่าเฉลี่ยก่อนแช่น้ำต้องมากกว่าศูนย์")
    ts = ((avg_a - avg_b) / avg_b) * 100.0
    return avg_b, avg_a, ts


def calc_statistics(values: list) -> dict:
    """คำนวณ mean, sd, cv, min, max"""
    arr = np.array([v for v in values if v is not None and not np.isnan(v)])
    if len(arr) == 0:
        return {"n": 0, "mean": 0, "sd": 0, "cv": 0, "min": 0, "max": 0}
    m = np.mean(arr)
    s = np.std(arr, ddof=1) if len(arr) > 1 else 0.0
    return {
        "n": len(arr),
        "mean": m,
        "sd": s,
        "cv": (s / m * 100) if m != 0 else 0,
        "min": np.min(arr),
        "max": np.max(arr),
    }

# =============================================================
# Helper: init session state
# =============================================================
def ss_get(key, default=0.0):
    return st.session_state.get(key, default)

def store_result(key, value):
    st.session_state[key] = value

# =============================================================
# Sidebar
# =============================================================
with st.sidebar:
    st.markdown("### ⚙️ Settings")

    st.markdown("---")
    st.markdown("#### 📂 Save / Load Parameters")

    uploaded_json = st.file_uploader("Upload JSON", type=["json"], label_visibility="collapsed")
    if uploaded_json is not None:
        try:
            loaded_data = json.load(uploaded_json)
            file_id = f"{uploaded_json.name}_{uploaded_json.size}"
            if ss_get("last_uploaded_file", "") != file_id:
                store_result("last_uploaded_file", file_id)
                for key, value in loaded_data.items():
                    st.session_state[key] = value
                st.success("✅ โหลดสำเร็จ")
                st.rerun()
        except Exception as e:
            st.error(f"❌ อ่านไฟล์ไม่ได้: {e}")

    # Export JSON
    def build_export_data():
        ns = int(ss_get("num_samples", 1))
        data = {"num_samples": ns}
        for i in range(ns):
            for k in ["L", "b", "t", "Fmax"]:
                data[f"{k}_{i}"] = ss_get(f"{k}_{i}", 0.0)
        for k in ["L_moe", "b_moe", "t_moe"]:
            data[k] = ss_get(k, 0.0)
        ns_ts = int(ss_get("num_ts_samples", 1))
        data["num_ts_samples"] = ns_ts
        for j in range(ns_ts):
            for side in range(1, 5):
                data[f"ts_b{side}_{j}"] = ss_get(f"ts_b{side}_{j}", 0.0)
                data[f"ts_a{side}_{j}"] = ss_get(f"ts_a{side}_{j}", 0.0)
        return data

    json_str = json.dumps(build_export_data(), ensure_ascii=False, indent=2)
    st.download_button("💾 Download JSON", data=json_str,
                       file_name="particleboard_params.json", mime="application/json")

    st.markdown("---")
    st.markdown("#### 📋 Standard Reference")
    standard = st.selectbox("มาตรฐานอ้างอิง", [
        "TIS 876-2547 (มอก.)",
        "JIS A 5908:2003",
        "EN 310:1993",
        "ASTM D1037",
        "Other / ไม่ระบุ"
    ], key="standard_ref")

    st.markdown("---")
    st.markdown(
        "<div style='text-align:center; font-size:12px; opacity:0.6; padding-top:20px;'>"
        "Particleboard Test v2.0</div>",
        unsafe_allow_html=True
    )

# =============================================================
# Header
# =============================================================
st.markdown("""
<div class="header-banner">
    <h1>🧪 Particleboard Bending Test</h1>
    <p>MOR • MOE (Elastic Region) • Thickness Swelling • Summary Report</p>
</div>
""", unsafe_allow_html=True)

# =============================================================
# Tabs
# =============================================================
tab_mor, tab_moe, tab_ts, tab_summary, tab_export = st.tabs([
    "① MOR", "② MOE + Graph", "③ Thickness Swelling", "④ Summary", "⑤ Export"
])

# =============================================================
# TAB 1: MOR
# =============================================================
with tab_mor:
    st.markdown('<div class="info-card">'
                '<b>Modulus of Rupture (MOR)</b> — Three-point bending test<br>'
                'MOR = 3·F<sub>max</sub>·L / (2·b·t²) &nbsp; [MPa = N/mm²]<br>'
                f'<small>📋 Standard: {standard}</small>'
                '</div>', unsafe_allow_html=True)

    num_samples = st.selectbox("จำนวนตัวอย่าง", [1, 2, 3, 4],
                               index=[1,2,3,4].index(int(ss_get("num_samples", 1))),
                               key="num_samples")

    mor_results = []

    for i in range(num_samples):
        with st.expander(f"📐 ตัวอย่างที่ {i+1}", expanded=(i == 0)):
            c1, c2, c3, c4 = st.columns(4)
            with c1:
                L = st.number_input("L (mm)", value=ss_get(f"L_{i}"), key=f"L_{i}",
                                    min_value=0.0, format="%.2f")
            with c2:
                b = st.number_input("b (mm)", value=ss_get(f"b_{i}"), key=f"b_{i}",
                                    min_value=0.0, format="%.2f")
            with c3:
                t = st.number_input("t (mm)", value=ss_get(f"t_{i}"), key=f"t_{i}",
                                    min_value=0.0, format="%.2f")
            with c4:
                Fmax_kg = st.number_input("F_max (kg)", value=ss_get(f"Fmax_{i}"), key=f"Fmax_{i}",
                                          min_value=0.0, format="%.3f")

            if L > 0 and b > 0 and t > 0 and Fmax_kg > 0:
                Fmax_N = Fmax_kg * 9.80665
                mor = calc_MOR(Fmax_N, L, b, t)
                mor_results.append(mor)
                store_result(f"mor_result_{i}", mor)
                st.markdown(f'<div class="result-card">'
                            f'<b>MOR ตัวอย่างที่ {i+1}</b> = <span style="font-size:22px;">'
                            f'{mor:.2f}</span> MPa &nbsp; '
                            f'<small>(F={Fmax_N:.2f} N)</small></div>',
                            unsafe_allow_html=True)
            else:
                mor_results.append(None)
                if any(v > 0 for v in [L, b, t, Fmax_kg]):
                    st.warning("กรุณาใส่ค่า L, b, t, Fmax ให้ครบและมากกว่า 0")

    # Summary
    valid_mor = [v for v in mor_results if v is not None]
    if len(valid_mor) >= 2:
        stats = calc_statistics(valid_mor)
        st.markdown("---")
        st.markdown("##### 📊 MOR Summary")
        mc1, mc2, mc3, mc4 = st.columns(4)
        with mc1:
            st.markdown(f'<div class="metric-box"><div class="label">Mean</div>'
                        f'<div class="value">{stats["mean"]:.2f}</div>'
                        f'<div class="unit">MPa</div></div>', unsafe_allow_html=True)
        with mc2:
            st.markdown(f'<div class="metric-box"><div class="label">Std Dev</div>'
                        f'<div class="value">{stats["sd"]:.2f}</div>'
                        f'<div class="unit">MPa</div></div>', unsafe_allow_html=True)
        with mc3:
            st.markdown(f'<div class="metric-box"><div class="label">CV</div>'
                        f'<div class="value">{stats["cv"]:.1f}</div>'
                        f'<div class="unit">%</div></div>', unsafe_allow_html=True)
        with mc4:
            st.markdown(f'<div class="metric-box"><div class="label">N</div>'
                        f'<div class="value">{stats["n"]}</div>'
                        f'<div class="unit">samples</div></div>', unsafe_allow_html=True)
        store_result("mor_stats", stats)
    elif len(valid_mor) == 1:
        store_result("mor_stats", {"n": 1, "mean": valid_mor[0], "sd": 0, "cv": 0,
                                   "min": valid_mor[0], "max": valid_mor[0]})

# =============================================================
# TAB 2: MOE + Load-Deflection
# =============================================================
with tab_moe:
    st.markdown('<div class="info-card">'
                '<b>Modulus of Elasticity (MOE)</b> — จาก Load-Deflection curve<br>'
                'คำนวณ Slope (N/mm) ของ elastic region ด้วย linear regression<br>'
                '<small>เลือกช่วง data points ที่เป็น linear (elastic) เพื่อ fit เส้นตรง<br>'
                'หากต้องการค่า MOE (MPa) ให้เปิด "คำนวณ MOE" แล้วใส่ geometry</small>'
                '</div>', unsafe_allow_html=True)

    # Template download
    template = pd.DataFrame({"Load (kg)": [0, 5, 10, 15, 20, 25],
                             "Deflection (mm)": [0.0, 0.8, 1.5, 2.3, 3.5, 5.0]})
    buf_tpl = io.BytesIO()
    template.to_excel(buf_tpl, index=False)
    buf_tpl.seek(0)
    st.download_button("📥 Download Template Excel", data=buf_tpl,
                       file_name="load_deflection_template.xlsx",
                       mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")

    uploaded_xl = st.file_uploader("อัปโหลดไฟล์ Excel (Load-Deflection)", type=["xlsx"])

    if uploaded_xl:
        try:
            df = pd.read_excel(uploaded_xl)
            if "Load (kg)" not in df.columns or "Deflection (mm)" not in df.columns:
                st.error("❌ ไฟล์ต้องมีคอลัมน์ 'Load (kg)' และ 'Deflection (mm)'")
            else:
                df["Load (N)"] = df["Load (kg)"] * 9.80665

                # แสดงตารางข้อมูลพร้อม index
                df_display = df.copy()
                df_display.insert(0, "Point #", range(len(df)))
                st.dataframe(df_display, use_container_width=True, height=220)

                load_N = df["Load (N)"].values
                defl = df["Deflection (mm)"].values
                n_pts = len(df)

                # --- เลือกช่วง Elastic Region ---
                st.markdown("##### 🎯 เลือกช่วง Elastic Region")
                st.markdown('<div class="warn-card">'
                            'เลือก <b>จุดเริ่มต้น</b> และ <b>จุดสิ้นสุด</b> ของช่วงที่กราฟเป็นเส้นตรง (elastic) '
                            'โดยดูจาก Point # ในตาราง'
                            '</div>', unsafe_allow_html=True)

                pc1, pc2 = st.columns(2)
                with pc1:
                    # Auto-detect: เริ่มจาก point ที่ 2 ของข้อมูลที่ load > 0
                    default_lo = 1 if n_pts > 2 else 0
                    idx_lo = st.number_input(
                        "Point # เริ่มต้น (Start)",
                        min_value=0, max_value=n_pts - 2,
                        value=int(ss_get("el_idx_lo", default_lo)),
                        step=1, key="el_idx_lo")
                with pc2:
                    # Auto-detect: ประมาณ 40-60% ของจำนวน point
                    default_hi = min(max(n_pts // 2, idx_lo + 1), n_pts - 1)
                    idx_hi = st.number_input(
                        "Point # สิ้นสุด (End)",
                        min_value=idx_lo + 1, max_value=n_pts - 1,
                        value=int(ss_get("el_idx_hi", default_hi)),
                        step=1, key="el_idx_hi")

                # --- คำนวณ Slope ---
                try:
                    slope, intercept, r_sq = calc_slope_elastic(load_N, defl, idx_lo, idx_hi)

                    store_result("moe_slope", slope)
                    store_result("moe_intercept", intercept)
                    store_result("moe_r2", r_sq)

                    # แสดงผล Slope
                    st.markdown(f'<div class="result-card">'
                                f'<b>Slope (Elastic Region)</b> &nbsp; '
                                f'Point #{idx_lo} → #{idx_hi} '
                                f'({idx_hi - idx_lo + 1} points)<br>'
                                f'<span style="font-size:26px;">{slope:.2f}</span> N/mm'
                                f'&nbsp;&nbsp;&nbsp;'
                                f'<span style="font-size:16px; color:#64748b;">'
                                f'R² = {r_sq:.4f}</span></div>',
                                unsafe_allow_html=True)

                    # --- Optional: คำนวณ MOE (MPa) ---
                    with st.expander("📐 คำนวณ MOE (MPa) — ต้องใส่ Geometry", expanded=False):
                        st.markdown("MOE = Slope × L³ / (4·b·t³)")
                        gc1, gc2, gc3 = st.columns(3)
                        with gc1:
                            L_moe = st.number_input("L (mm)", value=ss_get("L_moe", ss_get("L_0")),
                                                    key="L_moe", min_value=0.0, format="%.2f")
                        with gc2:
                            b_moe = st.number_input("b (mm)", value=ss_get("b_moe", ss_get("b_0")),
                                                    key="b_moe", min_value=0.0, format="%.2f")
                        with gc3:
                            t_moe = st.number_input("t (mm)", value=ss_get("t_moe", ss_get("t_0")),
                                                    key="t_moe", min_value=0.0, format="%.2f")

                        if L_moe > 0 and b_moe > 0 and t_moe > 0:
                            moe_val = slope_to_MOE(slope, L_moe, b_moe, t_moe)
                            store_result("moe_elastic", moe_val)
                            st.markdown(f'<div class="result-card">'
                                        f'<b>MOE</b> = '
                                        f'<span style="font-size:24px;">{moe_val:.2f}</span> MPa'
                                        f'</div>', unsafe_allow_html=True)
                        else:
                            st.info("ใส่ L, b, t ให้ครบเพื่อคำนวณ MOE (MPa)")

                    # --- กราฟ ---
                    fig, ax = plt.subplots(figsize=(9, 5.5))
                    fig.patch.set_facecolor('#fafbfe')
                    ax.set_facecolor('#fafbfe')

                    # All data points
                    ax.plot(defl, load_N, 'o-', color='#94a3b8', linewidth=1.5,
                            markersize=5, label='Test Data', zorder=2)

                    # Elastic region highlight
                    el_defl = defl[idx_lo:idx_hi+1]
                    el_load = load_N[idx_lo:idx_hi+1]
                    ax.plot(el_defl, el_load, 'o', color='#10b981', markersize=10,
                            markeredgecolor='white', markeredgewidth=1.5,
                            label=f'Elastic Region (#{idx_lo}–#{idx_hi})',
                            zorder=5)

                    # Linear fit line — ต่อเส้นให้ยาวพอเห็น
                    x_min_fit = max(0, defl[idx_lo] - (defl[idx_hi] - defl[idx_lo]) * 0.3)
                    x_max_fit = defl[idx_hi] + (defl[idx_hi] - defl[idx_lo]) * 0.5
                    x_fit = np.linspace(x_min_fit, x_max_fit, 100)
                    y_fit = slope * x_fit + intercept
                    ax.plot(x_fit, y_fit, '--', color='#ef4444', linewidth=2,
                            label=f'Linear Fit (slope={slope:.1f} N/mm, R²={r_sq:.3f})',
                            zorder=3)

                    # Annotate points
                    for k in [idx_lo, idx_hi]:
                        ax.annotate(f'#{k}', xy=(defl[k], load_N[k]),
                                    xytext=(8, 10), textcoords='offset points',
                                    fontsize=10, fontweight='bold', color='#10b981',
                                    bbox=dict(boxstyle='round,pad=0.3', facecolor='white',
                                              edgecolor='#10b981', alpha=0.9))

                    ax.set_xlabel("Deflection (mm)", fontsize=13, fontweight='bold')
                    ax.set_ylabel("Load (N)", fontsize=13, fontweight='bold')
                    ax.set_title("Load–Deflection Curve", fontsize=16, fontweight='bold')
                    ax.legend(fontsize=9, loc='best',
                              framealpha=0.9, edgecolor='#e2e8f0')
                    ax.grid(True, alpha=0.25, linestyle='--')
                    ax.spines[['top', 'right']].set_visible(False)
                    plt.tight_layout()

                    st.pyplot(fig)
                    store_result("moe_fig", fig)

                except Exception as e:
                    st.error(f"❌ คำนวณ Slope ไม่ได้: {e}")

        except Exception as e:
            st.error(f"❌ อ่านไฟล์ Excel ไม่ได้: {e}")

# =============================================================
# TAB 3: Thickness Swelling
# =============================================================
with tab_ts:
    st.markdown('<div class="info-card">'
                '<b>Thickness Swelling (TS)</b> — วัดความหนา 4 ด้าน ก่อน/หลังแช่น้ำ<br>'
                'TS = (t<sub>after</sub> − t<sub>before</sub>) / t<sub>before</sub> × 100 &nbsp; [%]'
                '</div>', unsafe_allow_html=True)

    num_ts = st.selectbox("จำนวนตัวอย่าง TS", [1, 2, 3, 4],
                          index=[1,2,3,4].index(int(ss_get("num_ts_samples", 1))),
                          key="num_ts_samples")

    ts_results = []

    for j in range(num_ts):
        with st.expander(f"📐 TS ตัวอย่างที่ {j+1}", expanded=(j == 0)):
            st.markdown("**ก่อนแช่น้ำ (mm)**")
            bc1, bc2, bc3, bc4 = st.columns(4)
            before = []
            for side, col in enumerate([bc1, bc2, bc3, bc4], 1):
                with col:
                    v = st.number_input(f"ด้าน {side}", value=ss_get(f"ts_b{side}_{j}"),
                                        key=f"ts_b{side}_{j}", min_value=0.0, format="%.3f")
                    before.append(v)

            st.markdown("**หลังแช่น้ำ (mm)**")
            ac1, ac2, ac3, ac4 = st.columns(4)
            after = []
            for side, col in enumerate([ac1, ac2, ac3, ac4], 1):
                with col:
                    v = st.number_input(f"ด้าน {side} ", value=ss_get(f"ts_a{side}_{j}"),
                                        key=f"ts_a{side}_{j}", min_value=0.0, format="%.3f")
                    after.append(v)

            if all(v > 0 for v in before) and all(v > 0 for v in after):
                avg_b, avg_a, ts = calc_TS(before, after)
                ts_results.append(ts)
                store_result(f"ts_result_{j}", ts)
                store_result(f"ts_avg_before_{j}", avg_b)
                store_result(f"ts_avg_after_{j}", avg_a)
                st.markdown(f'<div class="result-card">'
                            f'<b>TS ตัวอย่างที่ {j+1}</b> = '
                            f'<span style="font-size:22px;">{ts:.2f}</span> %<br>'
                            f'<small>Avg before = {avg_b:.3f} mm &nbsp;|&nbsp; '
                            f'Avg after = {avg_a:.3f} mm</small></div>',
                            unsafe_allow_html=True)
            else:
                ts_results.append(None)

    valid_ts = [v for v in ts_results if v is not None]
    if len(valid_ts) >= 2:
        ts_stats = calc_statistics(valid_ts)
        st.markdown("---")
        st.markdown("##### 📊 TS Summary")
        tc1, tc2, tc3, tc4 = st.columns(4)
        with tc1:
            st.markdown(f'<div class="metric-box"><div class="label">Mean</div>'
                        f'<div class="value">{ts_stats["mean"]:.2f}</div>'
                        f'<div class="unit">%</div></div>', unsafe_allow_html=True)
        with tc2:
            st.markdown(f'<div class="metric-box"><div class="label">Std Dev</div>'
                        f'<div class="value">{ts_stats["sd"]:.2f}</div>'
                        f'<div class="unit">%</div></div>', unsafe_allow_html=True)
        with tc3:
            st.markdown(f'<div class="metric-box"><div class="label">CV</div>'
                        f'<div class="value">{ts_stats["cv"]:.1f}</div>'
                        f'<div class="unit">%</div></div>', unsafe_allow_html=True)
        with tc4:
            st.markdown(f'<div class="metric-box"><div class="label">N</div>'
                        f'<div class="value">{ts_stats["n"]}</div>'
                        f'<div class="unit">samples</div></div>', unsafe_allow_html=True)
        store_result("ts_stats", ts_stats)

# =============================================================
# TAB 4: Summary
# =============================================================
with tab_summary:
    st.markdown("### 📊 Summary Table")

    # MOR table
    mor_data = []
    for i in range(int(ss_get("num_samples", 1))):
        v = ss_get(f"mor_result_{i}", None)
        if v is not None and v > 0:
            mor_data.append({
                "Sample": f"#{i+1}",
                "L (mm)": ss_get(f"L_{i}"),
                "b (mm)": ss_get(f"b_{i}"),
                "t (mm)": ss_get(f"t_{i}"),
                "Fmax (kg)": ss_get(f"Fmax_{i}"),
                "MOR (MPa)": round(v, 2),
            })

    if mor_data:
        st.markdown("##### MOR Results")
        df_mor = pd.DataFrame(mor_data)
        st.dataframe(df_mor, use_container_width=True, hide_index=True)
        mor_s = ss_get("mor_stats", {})
        if mor_s:
            st.markdown(f'<div class="summary-card">'
                        f'<b>MOR:</b> Mean = {mor_s.get("mean",0):.2f} MPa &nbsp;|&nbsp; '
                        f'SD = {mor_s.get("sd",0):.2f} MPa &nbsp;|&nbsp; '
                        f'CV = {mor_s.get("cv",0):.1f}%</div>',
                        unsafe_allow_html=True)

    # MOE
    moe_slope = ss_get("moe_slope", None)
    if moe_slope:
        st.markdown("##### MOE Results")
        moe_el = ss_get("moe_elastic", None)
        moe_text = f'<b>Slope (Elastic):</b> {moe_slope:.2f} N/mm &nbsp;|&nbsp; R² = {ss_get("moe_r2", 0):.4f}'
        if moe_el:
            moe_text += f'<br><b>MOE:</b> {moe_el:.2f} MPa'
        st.markdown(f'<div class="summary-card">{moe_text}</div>', unsafe_allow_html=True)

    # TS table
    ts_data = []
    for j in range(int(ss_get("num_ts_samples", 1))):
        v = ss_get(f"ts_result_{j}", None)
        if v is not None:
            ts_data.append({
                "Sample": f"#{j+1}",
                "Avg Before (mm)": round(ss_get(f"ts_avg_before_{j}", 0), 3),
                "Avg After (mm)": round(ss_get(f"ts_avg_after_{j}", 0), 3),
                "TS (%)": round(v, 2),
            })

    if ts_data:
        st.markdown("##### Thickness Swelling Results")
        df_ts = pd.DataFrame(ts_data)
        st.dataframe(df_ts, use_container_width=True, hide_index=True)
        ts_s = ss_get("ts_stats", {})
        if ts_s and ts_s.get("n", 0) >= 2:
            st.markdown(f'<div class="summary-card">'
                        f'<b>TS:</b> Mean = {ts_s.get("mean",0):.2f}% &nbsp;|&nbsp; '
                        f'SD = {ts_s.get("sd",0):.2f}% &nbsp;|&nbsp; '
                        f'CV = {ts_s.get("cv",0):.1f}%</div>',
                        unsafe_allow_html=True)

    if not mor_data and not moe_slope and not ts_data:
        st.info("ยังไม่มีผลลัพธ์ — กรุณาใส่ข้อมูลใน Tab ① ② ③ ก่อน")

# =============================================================
# TAB 5: Export
# =============================================================
with tab_export:
    st.markdown("### 📤 Export Report")

    exp_col1, exp_col2 = st.columns(2)

    # --- Excel Export ---
    with exp_col1:
        st.markdown("##### 📊 Excel Report")
        if st.button("สร้างไฟล์ Excel", use_container_width=True):
            with io.BytesIO() as buf:
                with pd.ExcelWriter(buf, engine='openpyxl') as writer:
                    # MOR sheet
                    if mor_data:
                        df_m = pd.DataFrame(mor_data)
                        mor_s = ss_get("mor_stats", {})
                        if mor_s and mor_s.get("n", 0) >= 1:
                            summary_row = {"Sample": "Summary", "MOR (MPa)": round(mor_s["mean"], 2)}
                            df_m = pd.concat([df_m, pd.DataFrame([summary_row])], ignore_index=True)
                        df_m.to_excel(writer, sheet_name="MOR", index=False)

                    # TS sheet
                    if ts_data:
                        df_t = pd.DataFrame(ts_data)
                        df_t.to_excel(writer, sheet_name="TS", index=False)

                    # MOE sheet
                    moe_slope_v = ss_get("moe_slope", None)
                    if moe_slope_v:
                        moe_row = {
                            "Slope (N/mm)": round(moe_slope_v, 2),
                            "R²": round(ss_get("moe_r2", 0), 4),
                        }
                        moe_el_v = ss_get("moe_elastic", None)
                        if moe_el_v:
                            moe_row["MOE (MPa)"] = round(moe_el_v, 2)
                        df_moe = pd.DataFrame([moe_row])
                        df_moe.to_excel(writer, sheet_name="MOE", index=False)

                buf.seek(0)
                st.download_button("📥 Download Excel", data=buf.getvalue(),
                                   file_name="particleboard_report.xlsx",
                                   mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")

    # --- Word Export ---
    with exp_col2:
        st.markdown("##### 📄 Word Report")
        if st.button("สร้าง Word Report", use_container_width=True):
            try:
                doc = Document()

                # Title
                title_para = doc.add_heading("Particleboard Bending Test Report", level=1)
                doc.add_paragraph(f"Standard: {ss_get('standard_ref', 'N/A')}")
                doc.add_paragraph("")

                # MOR section
                if mor_data:
                    doc.add_heading("1. Modulus of Rupture (MOR)", level=2)
                    doc.add_paragraph("MOR = 3·Fmax·L / (2·b·t²)  [MPa]")

                    table = doc.add_table(rows=1, cols=6, style='Light Shading Accent 1')
                    table.alignment = WD_TABLE_ALIGNMENT.CENTER
                    headers = ["Sample", "L (mm)", "b (mm)", "t (mm)", "Fmax (kg)", "MOR (MPa)"]
                    for idx_h, h in enumerate(headers):
                        table.rows[0].cells[idx_h].text = h

                    for row_data in mor_data:
                        row = table.add_row()
                        row.cells[0].text = str(row_data["Sample"])
                        row.cells[1].text = f'{row_data["L (mm)"]:.2f}'
                        row.cells[2].text = f'{row_data["b (mm)"]:.2f}'
                        row.cells[3].text = f'{row_data["t (mm)"]:.2f}'
                        row.cells[4].text = f'{row_data["Fmax (kg)"]:.3f}'
                        row.cells[5].text = f'{row_data["MOR (MPa)"]:.2f}'

                    mor_s = ss_get("mor_stats", {})
                    if mor_s and mor_s.get("n", 0) >= 2:
                        doc.add_paragraph(
                            f'Mean = {mor_s["mean"]:.2f} MPa  |  '
                            f'SD = {mor_s["sd"]:.2f} MPa  |  '
                            f'CV = {mor_s["cv"]:.1f}%')

                # MOE section
                moe_slope_v = ss_get("moe_slope", None)
                if moe_slope_v:
                    doc.add_heading("2. Modulus of Elasticity (MOE)", level=2)
                    doc.add_paragraph(
                        f'Slope (Elastic Region) = {moe_slope_v:.2f} N/mm  |  '
                        f'R² = {ss_get("moe_r2", 0):.4f}')
                    moe_el_v = ss_get("moe_elastic", None)
                    if moe_el_v:
                        doc.add_paragraph(f'MOE = {moe_el_v:.2f} MPa')

                # TS section
                if ts_data:
                    doc.add_heading("3. Thickness Swelling (TS)", level=2)
                    table2 = doc.add_table(rows=1, cols=4, style='Light Shading Accent 1')
                    table2.alignment = WD_TABLE_ALIGNMENT.CENTER
                    for idx_h, h in enumerate(["Sample", "Avg Before (mm)",
                                               "Avg After (mm)", "TS (%)"]):
                        table2.rows[0].cells[idx_h].text = h
                    for row_data in ts_data:
                        row = table2.add_row()
                        row.cells[0].text = str(row_data["Sample"])
                        row.cells[1].text = f'{row_data["Avg Before (mm)"]:.3f}'
                        row.cells[2].text = f'{row_data["Avg After (mm)"]:.3f}'
                        row.cells[3].text = f'{row_data["TS (%)"]:.2f}'

                    ts_s = ss_get("ts_stats", {})
                    if ts_s and ts_s.get("n", 0) >= 2:
                        doc.add_paragraph(
                            f'Mean = {ts_s["mean"]:.2f}%  |  '
                            f'SD = {ts_s["sd"]:.2f}%  |  '
                            f'CV = {ts_s["cv"]:.1f}%')

                # Footer
                doc.add_paragraph("")
                footer_p = doc.add_paragraph()
                footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = footer_p.add_run(
                    "พัฒนาโดย: รศ.ดร.อิทธิพล มีผล\n"
                    "ภาควิชาครุศาสตร์โยธา มจพ.")
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(100, 100, 100)

                buf_doc = io.BytesIO()
                doc.save(buf_doc)
                buf_doc.seek(0)

                st.download_button("📥 Download Word", data=buf_doc.getvalue(),
                                   file_name="particleboard_report.docx",
                                   mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

            except Exception as e:
                st.error(f"❌ ไม่สามารถสร้าง Word report ได้: {e}")

# =============================================================
# Footer
# =============================================================
st.markdown("""
<div class="footer">
    <b>พัฒนาโดย:</b> รศ.ดร.อิทธิพล มีผล<br>
    ภาควิชาครุศาสตร์โยธา มหาวิทยาลัยเทคโนโลยีพระจอมเกล้าพระนครเหนือ (มจพ.)
</div>
""", unsafe_allow_html=True)
